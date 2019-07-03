import os
import datetime
import docx
import re
import pytest
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from selenium import webdriver
from dotenv import load_dotenv

load_dotenv()

### Data collecting from LinkedIn job page (Attributed from git username: lic9)

# download chromedriver from http://chromedriver.chromium.org/downloads
# and replace the path with where you download it



## MAKE SURE TO SET ENGLISH AS THE LANGUAGE (attributed from Prof. Rossetti, git username: s2t2)
# source: https://sqa.stackexchange.com/questions/9904/how-to-set-browser-locale-with-chromedriver-python#
options = webdriver.ChromeOptions()
#options.add_experimental_option("prefs", {"intl.accept_languages": "es"}) # THIS WORKS - GETS SPANISH
options.add_experimental_option("prefs", {"intl.accept_languages": "en,en_US"})

# download chromedriver from http://chromedriver.chromium.org/downloads
# and replace the path with where you download it
driver = webdriver.Chrome('E:\Python scripts\chromedriver', chrome_options=options)

print("------------------------------")
print("Welcome to Cover Letter Premium!")
print("------------------------------")
job_url = input("please input your URL: ")
# Please put in a market reserach related job or intern.


### Validation URL (attributed from https://stackoverflow.com/questions/827557/how-do-you-validate-a-url-with-a-regular-expression-in-python)

def validate_url(url):
    regex = re.compile(
        r'^(?:http|ftp)s?://' # http:// or https://
        r'(?:(?:[A-Z0-9](?:[A-Z0-9-]{0,61}[A-Z0-9])?\.)+(?:[A-Z]{2,6}\.?|[A-Z0-9-]{2,}\.?)|' #domain...
        r'localhost|' #localhost...
        r'\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})' # ...or ip
        r'(?::\d+)?' # optional port
        r'(?:/?|[/?]\S+)$', re.IGNORECASE)
    return re.match(regex, url) is not None

while not validate_url(job_url):
    ## directly shut down the system
    print("------------------------------")
    print("Oh, expecting a valid URL. Please try again...")
    print("------------------------------")
    print('Shutting the program down...')
    exit()
    ## OR ask for input until the input is valid
    # job_url = input("Please input a valid URL: ")



print("------------------------------")
print("Getting the job information...")
print("------------------------------")
print("Based on the key words in the job description, here is your cover letter:")
print("------------------------------")


driver.get(job_url)
#driver.get('https://www.linkedin.com/jobs/view/1304372932/')


job_info =  driver.find_element_by_tag_name('title').get_attribute('text')
info_list = job_info.split('hiring')
rest_info = info_list[1].split('|')[0]
title = rest_info.split('in')[0]
address = rest_info.split('in')[1]

company = info_list[0]
description = driver.find_element_by_css_selector("div.description__text--rich").text

job_dict = {
	'title': title,
	'company': company,
	'address': address,
	'description': description
}
# driver.quit()
# print(job_dict)


### USER  Profile
user_name = os.environ.get("USER_name")
user_status = os.environ.get("USER_status")
user_major = os.environ.get("USER_major")
user_address = os.environ.get("USER_address")
user_phone = os.environ.get("USER_phone")
user_email = os.environ.get("USER_email")
user_school = os.environ.get("USER_school")
user_graduate = os.environ.get("USER_graduate")
# print(user_major)


### STARTING
# March, 30, 2019
today = datetime.date.today()
F_today = today.strftime("%B, %d, %Y")
# print(F_today)

# company_name =  "Phoenix Marketing International" 
company_name =  job_dict["company"]
company_address = job_dict["address"]
# print(company_name)

start = f""" 
{F_today}

{company_name}
{company_address}
"""
print(start)


### BOBY CONTENTS
# job_title = "Market research Intern" 
job_title =  job_dict["title"]

skill_1 = "analytical" 
skill_2 = "communication"  
skill_3 = "leadership"
skill_4 = "programming"
skill_5 = "time-management"
research = "past research experience,"
intro_p = f"I want to express my immense interest in the opportunity of the{job_title} at {company_name}. I am a {user_status.lower()} in the {user_major} at {user_school} graduating in {user_graduate}. For this{job_title} position, I am excited to apply my market research and consumer behavior knowledge to real-world situations and learn valuable experience from industry professionals. I believe that my qualifications and educational pursuits are a great fit with the kind of candidate you company is looking for. Given my passion and knowledge, I would be able to contribute to this role immediately."
analytical_research = "I know this position requires strong analytical and research ability, and my previous research experiences strongly supported my competence. I worked as a research assistant in two psychology labs at Willamette University for three years and actively engaged in both experimental design and data analysis processes. The complicated experimental design procedure improved my skills of thorough design and organization and attention to details. The data analysis processes, including data coding, cleaning, analysis, visualization, and interpretation, gave me great opportunities to apply my statistical knowledge and strengthen my analytical skills. From my current master program, I have learned market research methods and design thinking, using interviews, focus groups, and other qualitative methods to achieve human-centered designs. I am trained to apply psychological theories to the marketing industry and be familiar with qualitative research and analysis as well."

communication = "My communication skills gained good practice during my research experience. During the data collection, I have to understand the purpose of my supervisors and transmit correct information to experiment participants to gain expected and unbiased results. This training enhanced both my communication clarity and confidence."
global_leadership = "My prior experience of running an international student council gave me an excellent exercise of both my communication skills and leadership, and provided me a global perspective."
software = "I feel comfortable working with Microsoft Office applications, such as Excel and Power Point, and analytical software, such as SPSS, SYSTAT and R."
multitasking = "I also developed the ability to multitask through balancing my schedule of different research projects, leadership position, and coursework."



### Output
print("Dear Human Resource Manager:")
print("")
print(intro_p)
print("")


###  Conditional selection of key words from the job description ###
job_description = job_dict["description"]

skill = []

skill_body = str()
research_experience =str()

## Analytical/Resarch Paragraph

# if "analytical" in job_description:
#     skill.append(skill_1)  
#     research_experience = research
#     print(analytical_research)
#     print("")
# elif "research" in job_description:
#     skill.append(skill_1)  
#     research_experience = research
#     print(analytical_research)
#     print("")
# elif "quantitative" in job_description:
#     skill.append(skill_1) 
#     research_experience = research
#     print(analytical_research)
#     print("")
# elif "market research" in job_description:
#     skill = skill.append(skill_1) 
#     research_experience = research
#     print(analytical_research)
#     print("")
# else:
#     skill = skill
#     research_experience = research_experience

if  "analytical" in job_description or "research" in job_description or "quantitative" in job_description or "market research" in job_description:
    skill.append(skill_1) 
    research_experience = research
    print(analytical_research)
    print("")
else:
    skill_body = skill_body
    skill = skill

## Skill Paragraph
# Communication-related skills

# if "communication" in job_description:
#     skill_body = communication 
#     skill.append(skill_2) 
# elif "verbal" in job_description:
#     skill_body = communication 
#     skill.append(skill_2) 
# elif "people" in job_description:
#     skill_body = communication 
#     skill.append(skill_2) 
# else:
#     skill_body = skill_body
#     skill = skill

if __name__ == "__main__":
        
    def add_communication(skill):
        skill.append(skill_2)   
    if "communication" in job_description or "verbal" in job_description or "people" in job_description:
        add_communication(skill)
        skill_body = communication
    else:
        skill_body = skill_body
        skill = skill   

# Leadership or global perspective

# if "global"  in job_description:
#     skill_body = (skill_body + global_leadership)
#     skill.append(skill_3) 
# elif "leadership" in job_description:
#     skill_body = skill_body + global_leadership
#     skill.append(skill_3)  
# else:
#     skill_body = skill_body
#     skill = skill

def add_leadership_global(skill):
    skill.append(skill_3)

if "global"  in job_description or "leadership" in job_description:
    add_leadership_global(skill)
    skill_body = skill_body + global_leadership
else:
    skill_body = skill_body
    skill = skill


# Basic software skills

# if "software" in job_description:
#     skill_body = skill_body + software
#     skill.append(skill_4) 
# elif "Microsoft Office" in job_description:
#     skill_body = skill_body + software
#     skill.append(skill_4)
# elif "Excel" in job_description:
#     skill_body = skill_body + software
#     skill.append(skill_4)
# elif "PowerPoint" in job_description:
#     skill_body = skill_body + software
#     skill.append(skill_4)
# elif "Word" in job_description:
#     skill_body = skill_body + software
#     skill.append(skill_4)
# elif "SPSS" in job_description:
#     skill_body = skill_body + software
#     skill.append(skill_4)
# elif "R" in job_description:
#     skill_body = skill_body + software
#     skill.append(skill_4)
# elif "programming" in job_description:
#     skill_body = skill_body + software
#     skill.append(skill_4)
# else:
#     skill_body = skill_body
#     skill = skill

def add_software(skill):
    skill.append(skill_4)

if "software" in job_description or "Microsoft Office" in job_description or "Excel" in job_description or "PowerPoint" in job_description or "Word" in job_description or "SPSS" in job_description or "R" in job_description or "programming" in job_description:
    add_software(skill)
    skill_body = skill_body + software
else:
    skill_body = skill_body
    skill = skill


# Time management skill

# if "multitask"in job_description:
#     skill_body = skill_body + multitasking
#     skill.append(skill_5)
# elif "organizational" in job_description:
#     skill_body = skill_body + multitasking
#     skill.append(skill_5)
# elif "time-management" in job_description:
#     skill_body = skill_body + multitasking
#     skill.append(skill_5)
# elif "time management" in job_description:
#     skill_body = skill_body + multitasking
#     skill.append(skill_5)
# else:
#     skill_body = skill_body
#     skill = skill

def add_time_management(skill):
    skill.append(skill_5)

if "multitask" in job_description or "organize" in job_description or "organizational" in job_description or "time-management" in job_description or "time management" in job_description:
    add_time_management(skill)
    skill_body = skill_body + multitasking
else:
    skill_body = skill_body
    skill = skill

# print(skill)
# print(skill_body)
# breakpoint()


# Adding a comma butween each item (Attributed from: https://stackoverflow.com/questions/5920643/add-an-item-between-each-item-already-in-the-list)
sep = [', '] * len(skill)
skill_c = list(sum(zip(skill, sep), ())[:-1])
# def comma(skill, item):
#     result = [item] * (len(skill) * 2 - 1)
#     result[0::2] = skill
#     return result
# comma(skill, '-')

# Convert list to string (Attributed from: https://stackoverflow.com/questions/5618878/how-to-convert-list-to-string)
skill_s = ''.join(skill_c)

# print((skill_s))
# breakpoint()

ending_body = f"I feel strongly that my {skill_s} skills, {research_experience} and consumer psychology background will make me an excellent fit for the role of{job_title}. My global cultural perspective can also contribute to this position. I would welcome the opportunity to discuss the position in further interview. Thank you for your time and consideration."
   
print(skill_body)

print("")
print(ending_body)


## ENDING
end = f"""
Sincerely,
{user_name}
"""
print(end)


## Writing the output to a word document
### Attributed from: https://python-docx.readthedocs.io/en/latest/
doc = docx.Document()
# header = doc.add_heading(user_name, 0) 
header = doc.add_heading(user_name, 3)
header_contact = doc.add_heading(f"{user_address} | {user_phone} | {user_phone}", 8)
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
header_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

# header_style_font = header_style.font
# header_style_font.name = 'Times New Roman'
# font.header_style_font = Pt(8)
# header_font = header.font
# header_font.size = Pt(10)
# header_font.name = 'Times New Roman'
# header_font.alignment = WD_ALIGN_PARAGRAPH.CENTER

p1 = doc.add_paragraph(start)
p2 = doc.add_paragraph("Dear Human Resource Manager:")
p3 = doc.add_paragraph(intro_p)
p4 = doc.add_paragraph(analytical_research)
p5 = doc.add_paragraph(skill_body)
p6 = doc.add_paragraph(ending_body)
p7 = doc.add_paragraph(end)
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)
doc.save(f'{company_name}-cover letter.docx')


